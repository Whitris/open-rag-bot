import docx


def extract_text_from_docx(file_path: str) -> tuple[str, str]:
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    title = ""
    for p in doc.paragraphs:
        if p.style.name in ("Title", "Heading 1") and p.text.strip():
            title = p.text.strip()
            break
    if not title and paragraphs:
        title = paragraphs[0]
    return full_text, title
