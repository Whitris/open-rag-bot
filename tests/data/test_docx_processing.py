import docx
import pytest

from src.data.docx import extract_text_from_docx


@pytest.fixture
def sample_docx(tmp_path):
    path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_heading("Document Title", level=1)
    doc.add_paragraph("First paragraph.")
    doc.save(str(path))
    return str(path)


def test_extract_text_and_title_from_docx_heading(sample_docx):
    text, title = extract_text_from_docx(sample_docx)
    assert "Document Title" in text
    assert title == "Document Title"


def test_extract_text_and_title_from_docx_fallback(tmp_path):
    path = tmp_path / "plain.docx"
    doc = docx.Document()
    doc.add_paragraph("Fallback Title")
    doc.add_paragraph("Another paragraph.")
    doc.save(str(path))
    text, title = extract_text_from_docx(str(path))
    assert "Fallback Title" in text
    assert title == "Fallback Title"
